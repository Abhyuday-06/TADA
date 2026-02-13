"""
Export Module for DBMS Lab Exercises
Generates Word (DOCX) and PDF documents from assignment data.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ==========================================
#         WORD DOCUMENT GENERATION
# ==========================================

def generate_docx(
    user_config: dict,
    exercise_info: dict,
    setup_items: list[dict],
    query_items: list[dict],
    image_paths: list[str],
    output_path: str,
    font_name: str = "Calibri"
):
    """
    Generate a Word document matching the sample format.
    
    Args:
        user_config: {"name", "regno", "slot", "classNo", "faculty"}
        exercise_info: {"exercise_number", "exercise_title"}
        setup_items: [{"q": desc, "sql": sql, "result": result_text}, ...]
        query_items: [{"q": question, "sql": sql, "result": result_text}, ...]
        image_paths: List of screenshot image paths (one per query item only)
        output_path: Where to save the .docx file
        font_name: Font to use in the document (default: Calibri)
    """
    doc = Document()
    
    # -- Set default font --
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(11)
    
    # ===========================
    # HEADER SECTION
    # ===========================
    title = doc.add_paragraph('Database Systems Lab')
    title.style = doc.styles['Title']
    
    # Metadata fields
    meta_fields = [
        f"Experiment Name/No: {exercise_info['exercise_title']}(Ex. {exercise_info['exercise_number']})",
        f"Name: {user_config.get('name', 'Student')}",
        f"Registration Number: {user_config.get('regno', 'XXXXXXXXXX')}",
        f"Lab Slot: {user_config.get('slot', 'L00+00')}",
        f"Class Number: {user_config.get('classNo', '0000000000000')}",
    ]
    
    for field in meta_fields:
        p = doc.add_paragraph(field)
        p.style = doc.styles['Heading 2']
    
    doc.add_paragraph('')  # Spacer
    
    import re
    
    # ===========================
    # SETUP SECTION (CREATE + INSERT) - SQL only, no screenshots
    # ===========================
    current_table_name = None
    
    for item in setup_items:
        q_text = item.get('q', '')
        sql = item.get('sql', '')
        
        # Add table name heading for CREATE TABLE
        if sql.strip().lower().startswith('create'):
            match = re.search(r'CREATE\s+TABLE\s+(\S+)', sql, re.IGNORECASE)
            if match:
                table_name = match.group(1)
                if table_name != current_table_name:
                    current_table_name = table_name
                    heading = doc.add_paragraph(f'{table_name} table')
                    heading.style = doc.styles['Heading 2']
        
        # Description
        q_para = doc.add_paragraph(q_text)
        q_para.style = doc.styles['Normal']
        
        # SQL in monospace (no screenshot for setup)
        sql_para = doc.add_paragraph()
        run = sql_para.add_run(sql)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        
        doc.add_paragraph('')  # Spacer
    
    # ===========================
    # PRACTICE QUERIES - question + SQL + screenshot
    # ===========================
    if query_items:
        heading = doc.add_paragraph('Practice Queries')
        heading.style = doc.styles['Heading 1']
    
    for i, item in enumerate(query_items):
        q_text = item.get('q', f'Query {i+1}')
        sql = item.get('sql', '')
        
        # Question text
        q_para = doc.add_paragraph(q_text)
        q_para.style = doc.styles['Normal']
        
        # SQL in monospace
        sql_para = doc.add_paragraph()
        run = sql_para.add_run(sql)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        
        # Screenshot
        if i < len(image_paths) and os.path.exists(image_paths[i]):
            try:
                doc.add_picture(image_paths[i], width=Inches(6.0))
            except Exception as e:
                p = doc.add_paragraph(f'[Image not available: {e}]')
        
        doc.add_paragraph('')  # Spacer
    
    # Save the document
    doc.save(output_path)
    print(f"  [Export] Word document saved: {output_path}")
    return output_path


# ==========================================
#         PDF GENERATION
# ==========================================

def convert_to_pdf(docx_path: str, pdf_path: str = None) -> str:
    """
    Convert a Word document to PDF.
    Uses docx2pdf (requires Microsoft Word on Windows).
    
    Args:
        docx_path: Path to the .docx file
        pdf_path: Output PDF path (default: same name with .pdf extension)
    
    Returns:
        Path to the generated PDF
    """
    if pdf_path is None:
        pdf_path = docx_path.replace('.docx', '.pdf')
    
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        print(f"  [Export] PDF saved: {pdf_path}")
        return pdf_path
    except ImportError:
        print("  [!] docx2pdf not installed. Install with: pip install docx2pdf")
        print("  [!] Alternatively, open the .docx in Word and Save As PDF.")
        return None
    except Exception as e:
        print(f"  [!] PDF conversion failed: {e}")
        print("  [!] Make sure Microsoft Word is installed on your system.")
        print("  [!] You can manually open the .docx file and Save As PDF.")
        return None


# ==========================================
#           TEST / DEMO
# ==========================================
if __name__ == "__main__":
    # Quick test with mock data
    user_config = {
        "name": "Test User",
        "regno": "24BCE0000",
        "slot": "L19+20",
        "classNo": "2025260503021"
    }
    
    exercise_info = {
        "exercise_number": "4",
        "exercise_title": "SQL Operators"
    }
    
    setup_items = [
        {"q": "Create EMPLOYEE table", "sql": "CREATE TABLE test_employee (emp_id INT, emp_name VARCHAR(30))"},
        {"q": "Insert row 1", "sql": "INSERT INTO test_employee VALUES (1, 'Anand')"},
    ]
    
    query_items = [
        {"q": "A1. Display employee salary after adding a bonus of 5000.", "sql": "SELECT emp_name, salary + 5000 FROM test_employee"},
    ]
    
    output = generate_docx(user_config, exercise_info, setup_items, query_items, [], "test_output.docx")
    print(f"Test document created: {output}")
